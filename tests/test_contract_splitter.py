#!/usr/bin/env python3
"""
Test suite for contract_splitter package.
"""

import unittest
import tempfile
import os
from pathlib import Path
import sys

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from contract_splitter import ContractSplitter, split_document, flatten_sections
from contract_splitter.utils import (
    count_tokens, split_chinese_sentences, clean_text, 
    detect_heading_level, sliding_window_split
)


class TestUtils(unittest.TestCase):
    """Test utility functions."""
    
    def test_count_tokens_character(self):
        """Test character-based token counting."""
        text = "Hello world 你好世界"
        count = count_tokens(text, "character")
        self.assertEqual(count, len(text))
    
    def test_count_tokens_tiktoken(self):
        """Test tiktoken-based counting (if available)."""
        text = "Hello world"
        try:
            count = count_tokens(text, "tiktoken")
            self.assertIsInstance(count, int)
            self.assertGreater(count, 0)
        except ImportError:
            # tiktoken not available, should fall back to character count
            count = count_tokens(text, "tiktoken")
            self.assertEqual(count, len(text))
    
    def test_split_chinese_sentences(self):
        """Test Chinese sentence splitting."""
        text = "第一条规定了基本原则。第二条明确了双方责任！第三条约定了违约后果？"
        sentences = split_chinese_sentences(text)
        
        expected = [
            "第一条规定了基本原则",
            "第二条明确了双方责任",
            "第三条约定了违约后果"
        ]
        self.assertEqual(sentences, expected)
    
    def test_clean_text(self):
        """Test text cleaning."""
        text = "  Hello   world  \n\n  "
        cleaned = clean_text(text)
        self.assertEqual(cleaned, "Hello world")
    
    def test_detect_heading_level(self):
        """Test heading level detection."""
        test_cases = [
            ("第一章 总则", 1),
            ("第二节 基本原则", 2),
            ("一、合同目的", 3),
            ("（一）基本要求", 4),
            ("Chapter 1", 1),
            ("Section 2", 2),
            ("1. Introduction", 3),
            ("1.1 Overview", 4)
        ]
        
        for text, expected_level in test_cases:
            level = detect_heading_level(text)
            self.assertEqual(level, expected_level, f"Failed for: {text}")
    
    def test_sliding_window_split(self):
        """Test sliding window splitting."""
        text = "这是第一句。这是第二句。这是第三句。这是第四句。"
        chunks = sliding_window_split(text, max_tokens=10, overlap=3, by_sentence=True)
        
        self.assertGreater(len(chunks), 1)
        for chunk in chunks:
            self.assertLessEqual(len(chunk), 10)


class TestContractSplitter(unittest.TestCase):
    """Test main ContractSplitter class."""
    
    def setUp(self):
        """Set up test fixtures."""
        self.splitter = ContractSplitter(max_tokens=100, overlap=20)
    
    def test_initialization(self):
        """Test splitter initialization."""
        self.assertEqual(self.splitter.max_tokens, 100)
        self.assertEqual(self.splitter.overlap, 20)
        self.assertTrue(self.splitter.split_by_sentence)
        self.assertEqual(self.splitter.token_counter, "character")
    
    def test_unsupported_file_type(self):
        """Test handling of unsupported file types."""
        with self.assertRaises(ValueError):
            self.splitter.split("test.txt")
    
    def test_nonexistent_file(self):
        """Test handling of non-existent files."""
        with self.assertRaises(FileNotFoundError):
            self.splitter.split("nonexistent.pdf")
    
    def test_flatten_empty_sections(self):
        """Test flattening empty sections."""
        sections = []
        chunks = self.splitter.flatten(sections)
        self.assertEqual(chunks, [])
    
    def test_flatten_simple_sections(self):
        """Test flattening simple sections."""
        sections = [
            {
                "heading": "第一章",
                "content": "这是第一章的内容。",
                "level": 1,
                "subsections": []
            },
            {
                "heading": "第二章",
                "content": "这是第二章的内容。",
                "level": 1,
                "subsections": []
            }
        ]
        
        chunks = self.splitter.flatten(sections)
        self.assertEqual(len(chunks), 2)
        self.assertIn("第一章", chunks[0])
        self.assertIn("第二章", chunks[1])
    
    def test_flatten_hierarchical_sections(self):
        """Test flattening hierarchical sections."""
        sections = [
            {
                "heading": "第一章",
                "content": "章节内容",
                "level": 1,
                "subsections": [
                    {
                        "heading": "第一节",
                        "content": "节的内容",
                        "level": 2,
                        "subsections": []
                    }
                ]
            }
        ]
        
        chunks = self.splitter.flatten(sections)
        self.assertEqual(len(chunks), 2)
        self.assertIn("第一章", chunks[0])
        self.assertIn("第一章 > 第一节", chunks[1])


class TestDocumentCreation(unittest.TestCase):
    """Test document creation and processing."""
    
    def create_sample_docx(self):
        """Create a sample DOCX file for testing."""
        try:
            from docx import Document
            
            doc = Document()
            doc.add_heading('第一章 总则', level=1)
            doc.add_paragraph('本章规定了合同的基本原则。')
            
            doc.add_heading('第一条 合同目的', level=2)
            doc.add_paragraph('本合同是甲方与乙方就项目合作事宜达成的协议。')
            
            doc.add_heading('第二章 合作内容', level=1)
            doc.add_paragraph('本章详细说明了合作的具体内容。')
            
            # Save to temporary file
            temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
            doc.save(temp_file.name)
            return temp_file.name
            
        except ImportError:
            return None
    
    def test_docx_processing(self):
        """Test DOCX file processing if python-docx is available."""
        docx_file = self.create_sample_docx()
        
        if docx_file:
            try:
                splitter = ContractSplitter(max_tokens=200, overlap=50)
                sections = splitter.split(docx_file)
                
                self.assertGreater(len(sections), 0)
                
                # Check that we have hierarchical structure
                has_headings = any(section.get('heading') for section in sections)
                self.assertTrue(has_headings)
                
                # Test flattening
                chunks = splitter.flatten(sections)
                self.assertGreater(len(chunks), 0)
                
            finally:
                # Clean up
                os.unlink(docx_file)
        else:
            self.skipTest("python-docx not available")


class TestConvenienceFunctions(unittest.TestCase):
    """Test convenience functions."""
    
    def test_split_document_function(self):
        """Test split_document convenience function."""
        # This will fail with FileNotFoundError, but we test the function exists
        with self.assertRaises(FileNotFoundError):
            split_document("nonexistent.pdf")
    
    def test_flatten_sections_function(self):
        """Test flatten_sections convenience function."""
        sections = [
            {
                "heading": "Test",
                "content": "Content",
                "level": 1,
                "subsections": []
            }
        ]
        
        chunks = flatten_sections(sections)
        self.assertEqual(len(chunks), 1)
        self.assertIn("Test", chunks[0])


class TestIntegration(unittest.TestCase):
    """Integration tests."""
    
    def test_full_workflow(self):
        """Test complete workflow with sample data."""
        # Create sample hierarchical data
        sections = [
            {
                "heading": "第一章 总则",
                "content": "本章规定了合同的基本原则和适用范围。双方应严格遵守本合同的各项条款。",
                "level": 1,
                "subsections": [
                    {
                        "heading": "第一条 合同目的",
                        "content": "本合同是甲方与乙方就项目合作事宜达成的协议，旨在明确双方的权利和义务。",
                        "level": 2,
                        "subsections": []
                    }
                ]
            }
        ]
        
        # Test flattening
        splitter = ContractSplitter(max_tokens=50, overlap=10)
        chunks = splitter.flatten(sections)
        
        # Verify results
        self.assertGreater(len(chunks), 0)
        
        # Check that content is properly split
        total_content = "".join(chunks)
        self.assertIn("第一章", total_content)
        self.assertIn("第一条", total_content)
        
        # Verify chunk sizes
        for chunk in chunks:
            self.assertLessEqual(len(chunk), 100)  # Allow some flexibility


def run_tests():
    """Run all tests and save results."""
    # Create test suite
    test_suite = unittest.TestSuite()
    
    # Add test classes
    test_classes = [
        TestUtils,
        TestContractSplitter,
        TestDocumentCreation,
        TestConvenienceFunctions,
        TestIntegration
    ]
    
    for test_class in test_classes:
        tests = unittest.TestLoader().loadTestsFromTestCase(test_class)
        test_suite.addTests(tests)
    
    # Run tests
    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(test_suite)
    
    # Save results to output directory
    output_dir = Path("output")
    output_dir.mkdir(exist_ok=True)
    
    with open(output_dir / "test_results.txt", "w", encoding="utf-8") as f:
        f.write(f"Test Results\n")
        f.write(f"============\n\n")
        f.write(f"Tests run: {result.testsRun}\n")
        f.write(f"Failures: {len(result.failures)}\n")
        f.write(f"Errors: {len(result.errors)}\n")
        f.write(f"Success rate: {(result.testsRun - len(result.failures) - len(result.errors)) / result.testsRun * 100:.1f}%\n\n")
        
        if result.failures:
            f.write("Failures:\n")
            for test, traceback in result.failures:
                f.write(f"- {test}: {traceback}\n")
        
        if result.errors:
            f.write("Errors:\n")
            for test, traceback in result.errors:
                f.write(f"- {test}: {traceback}\n")
    
    return result


if __name__ == "__main__":
    print("Running contract_splitter test suite...")
    result = run_tests()
    
    if result.wasSuccessful():
        print("\n✓ All tests passed!")
    else:
        print(f"\n✗ {len(result.failures + result.errors)} test(s) failed")
        print("Check output/test_results.txt for details")
