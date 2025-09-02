#!/usr/bin/env python3
"""
调试表格结构
"""

import os
import sys
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from contract_splitter.docx_splitter import DocxSplitter
from docx import Document


def debug_table_structure():
    """调试表格结构"""
    print("🔍 调试表格结构")
    print("=" * 60)
    
    test_file = "output/【立项申请】首创证券新增代销机构广州农商行的立项申请.doc"
    
    if not os.path.exists(test_file):
        print(f"❌ 测试文件不存在: {test_file}")
        return
    
    # 创建DocxSplitter
    splitter = DocxSplitter(max_tokens=1000, overlap=100)
    
    # 先转换文件
    print("📄 转换.doc文件...")
    try:
        from contract_splitter.converter import DocumentConverter

        converter = DocumentConverter(cleanup_temp_files=False)  # 不清理，方便调试
        docx_path = converter.convert_to_docx(test_file)
        print(f"✅ 转换成功: {docx_path}")

        # 打开转换后的文档
        doc = Document(docx_path)
        
        print(f"\n📊 文档统计:")
        print(f"   段落数: {len(doc.paragraphs)}")
        print(f"   表格数: {len(doc.tables)}")
        
        # 分析每个表格
        for table_idx, table in enumerate(doc.tables):
            print(f"\n📋 表格 {table_idx + 1}:")
            print(f"   行数: {len(table.rows)}")
            print(f"   列数: {len(table.columns) if table.rows else 0}")
            
            # 分析每个cell
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    cell_length = len(cell_text)
                    
                    if cell_length > 0:
                        print(f"   Cell ({i+1},{j+1}): {cell_length} 字符")
                        if cell_length > 100:
                            preview = cell_text[:100].replace('\n', ' ')
                            print(f"      预览: {preview}...")
                        elif cell_length > 20:
                            preview = cell_text[:50].replace('\n', ' ')
                            print(f"      内容: {preview}")
                        else:
                            print(f"      内容: {cell_text}")
        
        # 测试嵌套表格提取
        print(f"\n🔍 测试嵌套表格提取:")
        for table_idx, table in enumerate(doc.tables):
            print(f"\n表格 {table_idx + 1}:")
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    try:
                        nested_content = splitter._extract_nested_tables_from_cell(cell)
                        if nested_content and len(nested_content) > 20:
                            print(f"   Cell ({i+1},{j+1}): 嵌套内容 {len(nested_content)} 字符")
                            if "广州金融控股" in nested_content:
                                print(f"      ✅ 包含股东信息")
                    except Exception as e:
                        print(f"   Cell ({i+1},{j+1}): 提取失败 - {e}")
        
    except Exception as e:
        print(f"❌ 处理失败: {e}")


def main():
    """主函数"""
    print("🚀 表格结构调试")
    print("=" * 80)
    
    debug_table_structure()
    
    print("\n" + "=" * 80)
    print("🎯 调试完成")


if __name__ == "__main__":
    main()
