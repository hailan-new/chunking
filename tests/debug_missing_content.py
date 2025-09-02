#!/usr/bin/env python3
"""
调试丢失的内容
"""

import os
import sys
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from contract_splitter.docx_splitter import DocxSplitter


def debug_missing_content():
    """调试丢失的内容"""
    print("🔍 调试丢失的内容")
    print("=" * 80)
    
    test_file = "output/【立项申请】首创证券新增代销机构广州农商行的立项申请.doc"
    
    if not os.path.exists(test_file):
        print(f"❌ 测试文件不存在: {test_file}")
        return
    
    # 创建DocxSplitter
    splitter = DocxSplitter(max_tokens=2000, overlap=200)
    
    print("📄 提取完整文档内容")
    try:
        from contract_splitter.converter import DocumentConverter
        
        converter = DocumentConverter(cleanup_temp_files=False)
        docx_path = converter.convert_to_docx(test_file)
        
        from docx import Document
        doc = Document(docx_path)
        
        # 提取所有文本内容
        all_text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                all_text += para.text.strip() + "\n"
        
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_content = splitter._extract_nested_tables_from_cell(cell)
                    if cell_content and cell_content.strip():
                        all_text += cell_content.strip() + "\n"
        
        print(f"📊 完整文档长度: {len(all_text)} 字符")
        
        # 搜索丢失的内容
        missing_phrases = [
            "广州农商行具有基金销售资格",
            "公司治理结构完善，内部控制有效",
            "组织架构完整，设有专门的产品销售业务团队",
            "配备相应的软硬件设施",
            "建立了产品销售业务的一整套流程和制度",
            "代销机构能够采取相应措施",
            "代销机构在销售过程及产品存续过程中",
            "负责产品销售业务的部门负责人",
            "计划产品销售业务与公司其他业务不存在利益冲突",
            "有健全的档案管理制度",
            "六、我司对于新增代销机构风险及管控措施"
        ]
        
        print(f"\n🔍 搜索丢失的内容:")
        for phrase in missing_phrases:
            if phrase in all_text:
                pos = all_text.find(phrase)
                context_start = max(0, pos - 50)
                context_end = min(len(all_text), pos + len(phrase) + 50)
                context = all_text[context_start:context_end]
                print(f"  ✅ 找到 '{phrase}' at position {pos}")
                print(f"     上下文: ...{context}...")
            else:
                print(f"  ❌ 未找到 '{phrase}'")
        
        # 现在测试分割
        print(f"\n📄 测试文档分割")
        sections = splitter.split(test_file)
        
        print(f"📊 分割结果: {len(sections)} 个sections")
        
        # 检查每个section的内容
        for i, section in enumerate(sections):
            print(f"\n📋 Section {i+1}:")
            if hasattr(section, 'title'):
                print(f"   标题: {section.title[:100]}...")
                print(f"   内容长度: {len(section.content)} 字符")
                section_text = section.title + "\n" + section.content
            else:
                # section是dict格式
                print(f"   标题: {section.get('title', 'N/A')[:100]}...")
                print(f"   内容长度: {len(section.get('content', ''))} 字符")
                section_text = section.get('title', '') + "\n" + section.get('content', '')

            # 检查是否包含丢失的内容
            found_phrases = []
            for phrase in missing_phrases:
                if phrase in section_text:
                    found_phrases.append(phrase)

            if found_phrases:
                print(f"   ✅ 包含: {found_phrases}")
            else:
                print(f"   ❌ 不包含任何丢失的内容")

        # 合并所有section内容检查
        all_section_content = ""
        for section in sections:
            if hasattr(section, 'title'):
                all_section_content += section.title + "\n" + section.content + "\n"
            else:
                all_section_content += section.get('title', '') + "\n" + section.get('content', '') + "\n"
        
        print(f"\n📊 所有sections合并长度: {len(all_section_content)} 字符")
        print(f"📊 原文档长度: {len(all_text)} 字符")
        print(f"📊 内容丢失: {len(all_text) - len(all_section_content)} 字符")
        
        print(f"\n🔍 在合并sections中搜索丢失内容:")
        for phrase in missing_phrases:
            if phrase in all_section_content:
                print(f"  ✅ 在sections中找到 '{phrase}'")
            else:
                print(f"  ❌ 在sections中未找到 '{phrase}'")
        
    except Exception as e:
        print(f"❌ 处理失败: {e}")
        import traceback
        traceback.print_exc()


def main():
    """主函数"""
    print("🚀 内容丢失调试")
    print("=" * 80)
    
    debug_missing_content()
    
    print("\n" + "=" * 80)
    print("🎯 调试完成")


if __name__ == "__main__":
    main()
