#!/usr/bin/env python3
"""
WPS文档处理模块
支持WPS文件的转换和处理
"""

import os
import logging
import shutil
import subprocess
import tempfile
from typing import Optional, List, Dict, Any
from pathlib import Path

# Windows COM支持
try:
    import win32com.client
    WIN32COM_AVAILABLE = True
except ImportError:
    WIN32COM_AVAILABLE = False

# WPS原生处理器支持
try:
    from .wps_native_processor import WPSNativeProcessor
    WPS_NATIVE_AVAILABLE = True
except ImportError:
    WPS_NATIVE_AVAILABLE = False

logger = logging.getLogger(__name__)


class WPSProcessor:
    """
    WPS文档处理器
    
    支持多种WPS文件转换方法：
    1. LibreOffice转换
    2. Windows COM接口转换
    3. 直接格式检测和重命名
    """
    
    def __init__(self, wps_api_key: Optional[str] = None):
        """
        初始化WPS处理器

        Args:
            wps_api_key: WPS开放平台API密钥（可选）
        """
        self._libreoffice_cmd = None  # 保存可用的LibreOffice命令
        self.wps_api_key = wps_api_key

        # 初始化WPS原生处理器
        if WPS_NATIVE_AVAILABLE:
            self.native_processor = WPSNativeProcessor(api_key=wps_api_key)
        else:
            self.native_processor = None

        self.available_converters = self._detect_available_converters()
        logger.info(f"可用的WPS转换器: {self.available_converters}")
    
    def _detect_available_converters(self) -> List[str]:
        """检测可用的转换器"""
        converters = []

        # 检查WPS原生处理器（优先级最高）
        if WPS_NATIVE_AVAILABLE and self.native_processor:
            converters.append("wps_native")

        # 检查LibreOffice
        if self._check_libreoffice():
            converters.append("libreoffice")

        # 检查Windows COM
        if WIN32COM_AVAILABLE and self._check_win32com():
            converters.append("win32com")

        # 直接格式检测总是可用
        converters.append("direct")

        return converters
    
    def _check_libreoffice(self) -> bool:
        """检查LibreOffice是否可用"""
        # 尝试多个可能的LibreOffice命令
        commands = ['libreoffice', 'soffice', '/opt/homebrew/bin/soffice']

        for cmd in commands:
            try:
                result = subprocess.run([cmd, '--version'],
                                      capture_output=True, text=True, timeout=5)
                if result.returncode == 0:
                    # 保存可用的命令
                    self._libreoffice_cmd = cmd
                    logger.info(f"LibreOffice found: {cmd}")
                    return True
            except:
                continue

        return False
    
    def _check_win32com(self) -> bool:
        """检查Windows COM是否可用"""
        if not WIN32COM_AVAILABLE:
            return False
        
        try:
            # 尝试创建Word应用程序对象
            word = win32com.client.Dispatch("Word.Application")
            word.Quit()
            return True
        except:
            return False
    
    def convert_to_docx(self, wps_file: str, output_dir: Optional[str] = None) -> Optional[str]:
        """
        将WPS文件转换为DOCX格式
        
        Args:
            wps_file: WPS文件路径
            output_dir: 输出目录，默认为WPS文件所在目录
            
        Returns:
            转换后的DOCX文件路径，失败返回None
        """
        if not os.path.exists(wps_file):
            logger.error(f"WPS文件不存在: {wps_file}")
            return None
        
        # 确定输出文件路径
        if output_dir is None:
            output_dir = os.path.dirname(wps_file)
        
        base_name = os.path.splitext(os.path.basename(wps_file))[0]
        output_file = os.path.join(output_dir, f"{base_name}.docx")
        
        # 按优先级尝试不同的转换器（WPS原生方案优先）
        for converter in self.available_converters:
            try:
                if converter == "wps_native":
                    native_result = self._convert_with_wps_native(wps_file, output_dir)
                    if native_result and os.path.exists(native_result):
                        logger.info(f"✓ WPS原生转换成功: {wps_file} -> {native_result}")
                        return native_result

                elif converter == "libreoffice":
                    if self._convert_with_libreoffice(wps_file, output_dir):
                        if os.path.exists(output_file):
                            logger.info(f"✓ LibreOffice转换成功: {wps_file} -> {output_file}")
                            return output_file

                elif converter == "win32com":
                    if self._convert_with_win32com(wps_file, output_file):
                        if os.path.exists(output_file):
                            logger.info(f"✓ Win32COM转换成功: {wps_file} -> {output_file}")
                            return output_file

                elif converter == "direct":
                    if self._convert_direct(wps_file, output_file):
                        if os.path.exists(output_file):
                            logger.info(f"✓ 直接转换成功: {wps_file} -> {output_file}")
                            return output_file

            except Exception as e:
                logger.warning(f"使用{converter}转换WPS失败: {e}")
                continue
        
        logger.error(f"所有转换方法都失败: {wps_file}")
        return None

    def _convert_with_wps_native(self, wps_file: str, output_dir: str) -> Optional[str]:
        """
        使用WPS原生方案转换文件

        Args:
            wps_file: WPS文件路径
            output_dir: 输出目录

        Returns:
            转换后的文件路径，失败返回None
        """
        if not self.native_processor:
            logger.warning("WPS原生处理器不可用")
            return None

        try:
            result = self.native_processor.convert_wps_to_docx(wps_file, output_dir)
            if result:
                logger.info(f"WPS原生转换成功: {result}")
                return result
            else:
                logger.warning("WPS原生转换失败")
                return None

        except Exception as e:
            logger.warning(f"WPS原生转换异常: {e}")
            return None

    def _convert_with_libreoffice(self, wps_file: str, output_dir: str) -> bool:
        """使用LibreOffice转换WPS文件"""
        if not self._libreoffice_cmd:
            logger.warning("LibreOffice命令未找到")
            return False

        try:
            # 尝试多种转换格式，以找到最佳的转换结果
            formats_to_try = [
                ('docx', 'Office Open XML Text'),
                ('rtf', 'Rich Text Format'),
                ('odt', 'writer8'),
                ('doc', 'MS Word 97')
            ]

            best_result = None
            best_size = 0

            for format_ext, filter_name in formats_to_try:
                try:
                    temp_output = os.path.join(output_dir, f"temp_{os.path.basename(wps_file).rsplit('.', 1)[0]}.{format_ext}")

                    cmd = [
                        self._libreoffice_cmd, '--headless', '--convert-to', format_ext,
                        '--outdir', output_dir, wps_file
                    ]

                    logger.info(f"尝试转换为{format_ext}: {' '.join(cmd)}")
                    result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)

                    if result.returncode == 0:
                        # 检查转换后的文件大小
                        converted_file = os.path.join(output_dir, f"{os.path.basename(wps_file).rsplit('.', 1)[0]}.{format_ext}")
                        if os.path.exists(converted_file):
                            file_size = os.path.getsize(converted_file)
                            logger.info(f"{format_ext}转换成功，文件大小: {file_size} bytes")

                            if file_size > best_size:
                                best_size = file_size
                                best_result = (format_ext, converted_file)

                except Exception as e:
                    logger.warning(f"{format_ext}转换失败: {e}")
                    continue

            if best_result:
                format_ext, best_file = best_result
                logger.info(f"选择最佳转换结果: {format_ext} (大小: {best_size} bytes)")

                # 优先尝试ODT或DOC格式，因为RTF可能有内容截断问题
                target_docx = os.path.join(output_dir, f"{os.path.basename(wps_file).rsplit('.', 1)[0]}.docx")

                if format_ext in ['odt', 'doc']:
                    # 将ODT/DOC文件标记为最佳源文件
                    best_source_marker = target_docx + f'.{format_ext}_source'
                    shutil.copy2(best_file, best_source_marker)
                    logger.info(f"保存{format_ext.upper()}源文件: {best_source_marker}")
                    return True
                elif format_ext == 'rtf':
                    # RTF作为备选方案
                    rtf_marker = target_docx + '.rtf_source'
                    shutil.copy2(best_file, rtf_marker)
                    logger.info(f"保存RTF源文件: {rtf_marker}")
                    return True

                # 如果最佳结果不是docx，尝试转换为docx
                elif format_ext != 'docx':
                    final_docx = os.path.join(output_dir, f"{os.path.basename(wps_file).rsplit('.', 1)[0]}.docx")
                    cmd = [
                        self._libreoffice_cmd, '--headless', '--convert-to', 'docx',
                        '--outdir', output_dir, best_file
                    ]
                    result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
                    if result.returncode == 0 and os.path.exists(final_docx):
                        logger.info(f"最终转换为docx成功: {final_docx}")
                        return True
                else:
                    return True

            logger.warning("所有转换格式都失败")
            return False

        except subprocess.TimeoutExpired:
            logger.warning("LibreOffice转换超时")
            return False
        except Exception as e:
            logger.warning(f"LibreOffice转换异常: {e}")
            return False
    
    def _convert_with_win32com(self, wps_file: str, output_file: str) -> bool:
        """使用Windows COM接口转换WPS文件"""
        if not WIN32COM_AVAILABLE:
            return False
        
        try:
            # 创建Word应用程序
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            try:
                # 打开WPS文件
                doc = word.Documents.Open(wps_file)
                
                # 保存为DOCX格式
                doc.SaveAs2(output_file, FileFormat=16)  # 16 = docx format
                doc.Close()
                
                return True
            
            finally:
                word.Quit()
        
        except Exception as e:
            logger.warning(f"Win32COM转换失败: {e}")
            return False
    
    def _convert_direct(self, wps_file: str, output_file: str) -> bool:
        """直接转换（某些WPS文件实际是DOCX格式）"""
        try:
            # 复制文件并重命名
            shutil.copy2(wps_file, output_file)
            
            # 尝试用python-docx验证
            try:
                import docx
                doc = docx.Document(output_file)
                # 如果能成功打开，说明转换成功
                return True
            except:
                # 如果打开失败，删除文件
                if os.path.exists(output_file):
                    os.remove(output_file)
                return False
        
        except Exception as e:
            logger.warning(f"直接转换失败: {e}")
            return False
    
    def extract_text(self, wps_file: str) -> Optional[str]:
        """
        从WPS文件提取文本
        
        Args:
            wps_file: WPS文件路径
            
        Returns:
            提取的文本内容，失败返回None
        """
        # 创建临时目录
        with tempfile.TemporaryDirectory() as temp_dir:
            # 转换为DOCX
            docx_file = self.convert_to_docx(wps_file, temp_dir)
            
            if docx_file and os.path.exists(docx_file):
                try:
                    # 使用docx库提取文本
                    import docx
                    doc = docx.Document(docx_file)
                    
                    text_parts = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text)
                    
                    # 处理表格
                    for table in doc.tables:
                        table_text = self._extract_table_text(table)
                        if table_text:
                            text_parts.append(table_text)
                    
                    full_text = '\n\n'.join(text_parts)
                    logger.info(f"✓ WPS文本提取成功: {len(full_text)} 字符")
                    return full_text
                
                except Exception as e:
                    logger.error(f"从转换后的DOCX提取文本失败: {e}")
        
        return None
    
    def _extract_table_text(self, table) -> str:
        """从表格提取文本"""
        table_text = []
        
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_text.append(cell_text)
            
            if any(row_text):  # 如果行中有内容
                table_text.append(" | ".join(row_text))
        
        return "\n".join(table_text)
    
    def is_wps_file(self, file_path: str) -> bool:
        """
        检测文件是否为WPS格式
        
        Args:
            file_path: 文件路径
            
        Returns:
            True表示是WPS文件
        """
        if not os.path.exists(file_path):
            return False
        
        # 检查文件扩展名
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.wps':
            return True
        
        # 可以添加更多的文件格式检测逻辑
        return False
    
    def get_conversion_info(self, wps_file: str) -> Dict[str, Any]:
        """
        获取WPS文件转换信息
        
        Args:
            wps_file: WPS文件路径
            
        Returns:
            转换信息字典
        """
        info = {
            'file_path': wps_file,
            'file_exists': os.path.exists(wps_file),
            'file_size': 0,
            'available_converters': self.available_converters,
            'recommended_converter': None
        }
        
        if info['file_exists']:
            info['file_size'] = os.path.getsize(wps_file)
            
            # 推荐转换器
            if "libreoffice" in self.available_converters:
                info['recommended_converter'] = "libreoffice"
            elif "win32com" in self.available_converters:
                info['recommended_converter'] = "win32com"
            else:
                info['recommended_converter'] = "direct"
        
        return info


# 便捷函数
def convert_wps_to_docx(wps_file: str, output_dir: Optional[str] = None) -> Optional[str]:
    """
    便捷函数：转换WPS文件为DOCX
    
    Args:
        wps_file: WPS文件路径
        output_dir: 输出目录
        
    Returns:
        转换后的DOCX文件路径
    """
    processor = WPSProcessor()
    return processor.convert_to_docx(wps_file, output_dir)


def extract_wps_text(wps_file: str) -> Optional[str]:
    """
    便捷函数：从WPS文件提取文本
    
    Args:
        wps_file: WPS文件路径
        
    Returns:
        提取的文本内容
    """
    processor = WPSProcessor()
    return processor.extract_text(wps_file)
