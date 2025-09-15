"""
DOCX document splitter for extracting hierarchical sections from Word documents.
"""

from typing import List, Dict, Any, Optional
import logging
import os
import subprocess
import tempfile
from pathlib import Path
from .base import BaseSplitter
from .utils import count_tokens, sliding_window_split, clean_text, detect_heading_level
from .legal_structure_detector import get_legal_detector

# Try to import textract for .doc file support
try:
    import textract
    TEXTRACT_AVAILABLE = True
except ImportError:
    TEXTRACT_AVAILABLE = False

# DocxSplitter focuses on DOCX and DOC files only

logger = logging.getLogger(__name__)


class DocxSplitter(BaseSplitter):
    """
    Splitter for .doc and .docx files using python-docx.
    
    Extracts document structure based on heading styles and paragraph content.
    Supports Chinese documents with proper text handling.
    """
    
    def __init__(self, max_tokens: int = 2000, overlap: int = 200,
                 split_by_sentence: bool = True, token_counter: str = "character",
                 chunking_strategy: str = "finest_granularity",
                 strict_max_tokens: bool = False,
                 use_llm_heading_detection: bool = False,
                 llm_config: dict = None):
        """
        Initialize DOCX splitter.

        Args:
            max_tokens: Maximum tokens per chunk
            overlap: Overlap length for sliding window
            split_by_sentence: Whether to split at sentence boundaries
            token_counter: Token counting method
            chunking_strategy: Chunking strategy for flatten operation
            strict_max_tokens: Whether to strictly enforce max_tokens limit
            use_llm_heading_detection: Whether to use LLM for heading detection
            llm_config: LLM configuration dict
        """
        super().__init__(max_tokens, overlap, split_by_sentence, token_counter,
                        chunking_strategy, strict_max_tokens)

        self.use_llm_heading_detection = use_llm_heading_detection
        self.llm_heading_detector = None

        # 初始化法律结构检测器
        self.structure_detector = get_legal_detector("legal")

        if use_llm_heading_detection:
            try:
                from .llm_heading_detector import create_llm_heading_detector
                # 如果没有提供llm_config，从全局配置获取
                if llm_config is None:
                    from .config import get_config
                    config = get_config()
                    llm_config = config.get_llm_config()

                self.llm_heading_detector = create_llm_heading_detector(llm_config)
                logger.info("LLM heading detection enabled")
            except Exception as e:
                logger.warning(f"Failed to initialize LLM heading detector: {e}")
                self.use_llm_heading_detection = False
        
    def split(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Split DOCX document into hierarchical sections.

        Args:
            file_path: Path to the DOCX file

        Returns:
            List of hierarchical section dictionaries
        """
        self.validate_file(file_path, ['.doc', '.docx'])

        # Try different approaches based on file type and format
        file_ext = os.path.splitext(file_path)[1].lower()

        if file_ext == '.doc':
            # Handle legacy .doc files
            return self._handle_legacy_doc(file_path)
        else:
            # Handle .docx files
            return self._handle_docx_file(file_path)

    def _handle_legacy_doc(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Handle legacy .doc files using automatic conversion and alternative methods.

        Args:
            file_path: Path to the .doc file

        Returns:
            List of hierarchical section dictionaries
        """
        # First, try automatic conversion to .docx
        try:
            logger.info(f"Attempting automatic conversion of .doc file: {file_path}")
            from .converter import DocumentConverter

            converter = DocumentConverter(cleanup_temp_files=True)
            converted_path = converter.convert_to_docx(file_path)

            logger.info(f"✓ Successfully converted to: {converted_path}")

            # Process the converted .docx file
            from docx import Document
            doc = Document(converted_path)
            elements = self._extract_elements(doc)
            elements = self._enhance_structure_detection(elements)
            sections = self._build_hierarchy(elements)
            result = self._apply_size_constraints(sections)

            # Cleanup will happen automatically when converter is destroyed
            return result

        except Exception as e1:
            logger.warning(f"Automatic conversion failed: {e1}")

        # Fallback to direct processing methods
        try:
            # Try python-docx first (sometimes works with .doc)
            from docx import Document
            doc = Document(file_path)
            elements = self._extract_elements(doc)
            elements = self._enhance_structure_detection(elements)
            sections = self._build_hierarchy(elements)
            return self._apply_size_constraints(sections)
        except Exception as e2:
            logger.warning(f"python-docx direct processing failed: {e2}")

            # Try alternative text extraction methods
            try:
                text = self._extract_with_antiword(file_path)
                return self._process_plain_text(text, file_path)
            except Exception as e3:
                logger.warning(f"antiword failed: {e3}")

            # Try alternative: use win32com (Windows only)
            if os.name == 'nt':
                try:
                    return self._convert_doc_with_win32com(file_path)
                except Exception as e4:
                    logger.warning(f"win32com failed: {e4}")

            # Try alternative: use pandoc
            try:
                text = self._extract_with_pandoc(file_path)
                return self._process_plain_text(text, file_path)
            except Exception as e5:
                logger.warning(f"pandoc failed: {e5}")

            # Last resort: provide helpful error message
            raise ValueError(f"Unable to process .doc file: {file_path}. "
                           f"This appears to be a legacy or non-standard .doc format. "
                           f"Automatic conversion failed: {e1}. "
                           f"Direct processing failed: {e2}. "
                           f"Install LibreOffice, pandoc, or convert manually to .docx format.")

    def _handle_docx_file(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Handle .docx files with fallback strategies.

        Args:
            file_path: Path to the .docx file

        Returns:
            List of hierarchical section dictionaries
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required for DOCX processing. "
                            "Install with: pip install python-docx")

        try:
            doc = Document(file_path)
            elements = self._extract_elements(doc)
            # Enhance structure detection for documents without clear headings
            elements = self._enhance_structure_detection(elements)
            sections = self._build_hierarchy(elements)
            return self._apply_size_constraints(sections)
        except Exception as e:
            logger.warning(f"Standard DOCX processing failed: {e}")

            # Try alternative: extract as plain text and process
            try:
                import docx2txt
                text = docx2txt.process(file_path)
                return self._process_plain_text(text, file_path)
            except ImportError:
                logger.warning("docx2txt not available for fallback")
            except Exception as e2:
                logger.warning(f"docx2txt fallback failed: {e2}")

            # Final fallback: try to read as zip and extract text manually
            try:
                return self._extract_docx_as_zip(file_path)
            except Exception as e3:
                logger.warning(f"ZIP extraction failed: {e3}")

            raise ValueError(f"Failed to open DOCX file: {e}. "
                           f"File may be corrupted or in an unsupported format.")
    
    def _extract_elements(self, doc) -> List[Dict[str, Any]]:
        """
        Extract all document elements (headings, paragraphs, and tables).

        Args:
            doc: python-docx Document object

        Returns:
            List of element dictionaries
        """
        elements = []

        # Process document in order (paragraphs and tables)
        for element in doc.element.body:
            if element.tag.endswith('p'):  # Paragraph
                # Find corresponding paragraph object
                for paragraph in doc.paragraphs:
                    if paragraph._element == element:
                        text = clean_text(paragraph.text)
                        if text:
                            elem = {
                                'text': text,
                                'style': paragraph.style.name if paragraph.style else 'Normal',
                                'is_heading': self._is_heading(paragraph),
                                'level': self._get_heading_level(paragraph),
                                'type': 'paragraph'
                            }
                            elements.append(elem)
                        break

            elif element.tag.endswith('tbl'):  # Table
                # Find corresponding table object
                for table in doc.tables:
                    if table._element == element:
                        # 按照用户方案：对每个table cell当作独立文档进行层级split
                        cell_elements = self._extract_single_table_cells(table)
                        elements.extend(cell_elements)
                        break

        # Fallback: if no elements found, use simple paragraph extraction
        if not elements:
            for paragraph in doc.paragraphs:
                text = clean_text(paragraph.text)
                if text:
                    element = {
                        'text': text,
                        'style': paragraph.style.name if paragraph.style else 'Normal',
                        'is_heading': self._is_heading(paragraph),
                        'level': self._get_heading_level(paragraph),
                        'type': 'paragraph'
                    }
                    elements.append(element)

        return elements
    
    def _is_heading(self, paragraph) -> bool:
        """
        Determine if a paragraph is a heading.
        
        Args:
            paragraph: python-docx Paragraph object
            
        Returns:
            True if paragraph is a heading
        """
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        
        # Check for heading styles
        if 'heading' in style_name:
            return True
        
        # Check for title styles
        if 'title' in style_name:
            return True
        
        # Check text patterns for Chinese/English headings
        text = paragraph.text.strip()
        if self._looks_like_heading(text):
            return True
        
        return False
    
    def _looks_like_heading(self, text: str) -> bool:
        """
        Check if text looks like a heading based on patterns.
        
        Args:
            text: Text to check
            
        Returns:
            True if text looks like a heading
        """
        import re
        
        # Chinese heading patterns
        chinese_patterns = [
            r'^第[一二三四五六七八九十\d]+章',  # Chapter
            r'^第[一二三四五六七八九十\d]+节',  # Section
            r'^[一二三四五六七八九十\d]+[、．.]',  # Numbered
            r'^（[一二三四五六七八九十\d]+）',  # Parenthetical
        ]
        
        # English heading patterns
        english_patterns = [
            r'^Chapter\s+\d+',
            r'^Section\s+\d+',
            r'^\d+\.?\s+',
            r'^\d+\.\d+\.?\s+',
        ]
        
        for pattern in chinese_patterns + english_patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return True
        
        # Check if text is short and might be a heading
        if len(text) < 100 and not text.endswith(('。', '.', '！', '!', '？', '?')):
            return True
        
        return False
    
    def _get_heading_level(self, paragraph) -> int:
        """
        Get heading level from paragraph.
        
        Args:
            paragraph: python-docx Paragraph object
            
        Returns:
            Heading level (1-6)
        """
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        
        # Extract level from style name
        if 'heading' in style_name:
            import re
            match = re.search(r'heading\s*(\d+)', style_name)
            if match:
                return int(match.group(1))
        
        # Use text pattern detection
        return detect_heading_level(paragraph.text)
    
    def _build_hierarchy(self, elements: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        Build hierarchical structure from flat list of elements.
        
        Args:
            elements: List of document elements
            
        Returns:
            List of hierarchical sections
        """
        sections = []
        current_section = None
        section_stack = []  # Stack to track nested sections
        
        for i, element in enumerate(elements):
            # Debug logging for table cell elements
            if element.get('source', '').startswith('table_cell'):
                logger.info(f"Processing table cell element {i}: is_heading={element.get('is_heading')}, level={element.get('level')}, text_length={len(element.get('text', ''))}")

            if element['is_heading']:
                # Create new section
                section = {
                    'heading': element['text'],
                    'content': element['text'],  # 标题文字也要包含在content中
                    'level': element['level'],
                    'subsections': []
                }

                # Determine where to place this section
                if not section_stack or element['level'] <= section_stack[-1]['level']:
                    # Pop sections until we find the right parent level
                    while section_stack and element['level'] <= section_stack[-1]['level']:
                        section_stack.pop()

                    if section_stack:
                        # Add as subsection to parent
                        section_stack[-1]['subsections'].append(section)
                    else:
                        # Add as top-level section
                        sections.append(section)
                else:
                    # Add as subsection to current section
                    if section_stack:
                        section_stack[-1]['subsections'].append(section)
                    else:
                        sections.append(section)

                section_stack.append(section)
                current_section = section
            else:
                # Add content to current section
                if element.get('source', '').startswith('table_cell'):
                    logger.info(f"Adding table cell content to current_section: {current_section['heading'][:50] if current_section else 'None'}...")

                if current_section is not None:
                    if current_section['content']:
                        current_section['content'] += '\n\n' + element['text']
                    else:
                        current_section['content'] = element['text']

                    if element.get('source', '').startswith('table_cell'):
                        logger.info(f"Table cell content added, new content length: {len(current_section['content'])}")
                else:
                    # No current section, create a default one
                    if not sections:
                        sections.append({
                            'heading': 'Document Content',
                            'content': element['text'],
                            'level': 1,
                            'subsections': []
                        })
                        current_section = sections[0]
                        section_stack = [current_section]
                    else:
                        # Add to last section
                        last_section = self._find_last_section(sections)
                        if last_section['content']:
                            last_section['content'] += '\n\n' + element['text']
                        else:
                            last_section['content'] = element['text']
        
        return sections

    def _enhance_structure_detection(self, elements: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        Enhance structure detection for documents without clear headings.

        Args:
            elements: List of document elements

        Returns:
            Enhanced elements with better structure detection
        """
        if not elements:
            return elements

        enhanced_elements = []

        for i, element in enumerate(elements):
            enhanced_element = element.copy()

            # If not already detected as heading, try additional detection
            if not enhanced_element.get('is_heading', False):
                text = enhanced_element.get('text', '').strip()

                # Use unified structure detector for heading detection
                if self.structure_detector.is_legal_heading(text):
                    enhanced_element['is_heading'] = True
                    # Determine level using structure detector
                    enhanced_element['level'] = self.structure_detector.get_heading_level(text)

            enhanced_elements.append(enhanced_element)

        return enhanced_elements

    def _extract_docx_as_zip(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract DOCX content by treating it as a ZIP file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            List of document elements
        """
        import zipfile
        import xml.etree.ElementTree as ET

        elements = []

        try:
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                # Read document.xml
                if 'word/document.xml' in zip_file.namelist():
                    with zip_file.open('word/document.xml') as doc_file:
                        tree = ET.parse(doc_file)
                        root = tree.getroot()

                        # Extract text from paragraphs
                        for para in root.iter():
                            if para.tag.endswith('}p'):  # paragraph
                                text_parts = []
                                for text_elem in para.iter():
                                    if text_elem.tag.endswith('}t'):  # text
                                        if text_elem.text:
                                            text_parts.append(text_elem.text)

                                if text_parts:
                                    text = ''.join(text_parts).strip()
                                    if text:
                                        # Basic heading detection
                                        is_heading = detect_heading_level(text) > 0
                                        level = detect_heading_level(text) if is_heading else 0

                                        elements.append({
                                            'text': text,
                                            'type': 'heading' if is_heading else 'paragraph',
                                            'is_heading': is_heading,
                                            'level': level,
                                            'source': 'zip_extraction'
                                        })

        except Exception as e:
            logger.warning(f"ZIP extraction failed: {e}")
            raise

        return elements

    def _extract_single_table_cells(self, table) -> List[Dict[str, Any]]:
        """
        Extract cells from a single table as separate document elements.

        Args:
            table: python-docx Table object

        Returns:
            List of table cell elements
        """
        elements = []

        try:
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        elements.append({
                            'text': cell_text,
                            'type': 'table_cell',
                            'is_heading': False,
                            'level': 0,
                            'source': f'table_cell_{row_idx}_{cell_idx}',
                            'table_info': {
                                'row_index': row_idx,
                                'cell_index': cell_idx
                            }
                        })

        except Exception as e:
            logger.warning(f"Single table cell extraction failed: {e}")

        return elements

    def _extract_table_cells_as_documents(self, doc) -> List[Dict[str, Any]]:
        """
        Extract table cells as separate document elements.

        Args:
            doc: python-docx Document object

        Returns:
            List of table cell elements
        """
        elements = []

        try:
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if cell_text:
                            elements.append({
                                'text': cell_text,
                                'type': 'table_cell',
                                'is_heading': False,
                                'level': 0,
                                'source': f'table_cell_{table_idx}_{row_idx}_{cell_idx}',
                                'table_info': {
                                    'table_index': table_idx,
                                    'row_index': row_idx,
                                    'cell_index': cell_idx
                                }
                            })

        except Exception as e:
            logger.warning(f"Table cell extraction failed: {e}")

        return elements

    def _apply_size_constraints(self, sections: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        Apply size constraints to sections.

        Args:
            sections: List of hierarchical sections

        Returns:
            Sections with size constraints applied
        """
        if not sections:
            return sections

        # If strict_max_tokens is enabled, ensure no section exceeds the limit
        if self.strict_max_tokens:
            return self._enforce_strict_size_limits(sections)

        return sections

    def _enforce_strict_size_limits(self, sections: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        Enforce strict size limits on sections.

        Args:
            sections: List of sections

        Returns:
            Sections with enforced size limits
        """
        processed_sections = []

        for section in sections:
            # Process subsections first
            if section.get('subsections'):
                section['subsections'] = self._enforce_strict_size_limits(section['subsections'])

            # Check section content size
            content = section.get('content', '')
            if content and count_tokens(content) > self.max_tokens:
                # Split large content
                split_content = sliding_window_split(
                    content,
                    self.max_tokens,
                    self.overlap,
                    by_sentence=self.split_by_sentence
                )

                # Create multiple sections from split content
                for i, chunk in enumerate(split_content):
                    new_section = section.copy()
                    new_section['content'] = chunk
                    new_section['subsections'] = [] if i > 0 else section.get('subsections', [])
                    if i > 0:
                        new_section['heading'] = f"{section['heading']} (Part {i+1})"
                    processed_sections.append(new_section)
            else:
                processed_sections.append(section)

        return processed_sections

    def extract_text(self, file_path: str) -> str:
        """
        Extract plain text from DOCX document efficiently

        Args:
            file_path: Path to the DOCX file

        Returns:
            Extracted plain text content
        """
        self.validate_file(file_path, ['.docx', '.doc'])

        try:
            # Try python-docx first for best results
            from docx import Document

            doc = Document(file_path)
            text_parts = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_texts.append(cell_text)
                    if row_texts:
                        text_parts.append(' | '.join(row_texts))

            return '\n'.join(text_parts)

        except ImportError:
            # Fallback to ZIP extraction
            try:
                return self._extract_text_as_zip(file_path)
            except Exception:
                # Final fallback: use split method
                logger.warning("No DOCX libraries available, using split method for text extraction")
                return super().extract_text(file_path)

    def _extract_text_as_zip(self, file_path: str) -> str:
        """
        Extract text from DOCX by treating it as ZIP file

        Args:
            file_path: Path to the DOCX file

        Returns:
            Extracted text content
        """
        import zipfile
        import xml.etree.ElementTree as ET

        text_parts = []

        with zipfile.ZipFile(file_path, 'r') as zip_file:
            # Read document.xml
            if 'word/document.xml' in zip_file.namelist():
                with zip_file.open('word/document.xml') as doc_file:
                    tree = ET.parse(doc_file)
                    root = tree.getroot()

                    # Extract text from paragraphs
                    for para in root.iter():
                        if para.tag.endswith('}p'):  # paragraph
                            para_texts = []
                            for text_elem in para.iter():
                                if text_elem.tag.endswith('}t'):  # text
                                    if text_elem.text:
                                        para_texts.append(text_elem.text)

                            if para_texts:
                                text = ''.join(para_texts).strip()
                                if text:
                                    text_parts.append(text)

        return '\n'.join(text_parts)
